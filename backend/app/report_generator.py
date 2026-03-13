import os
from docx import Document
from datetime import datetime

def generate_report(original_filename, parsed_content, result_json):
    """
    根据解析后的内容生成质检报告
    :param original_filename: 原始文件名（用于生成报告名）
    :param parsed_content: dedoc 解析出的文本内容（字符串）
    :param result_json: AI 质检结果
    :return: 报告文件路径
    """
    # 生成报告文件名
    base = os.path.splitext(original_filename)[0]
    report_path = f"{base}_质检报告_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

    # 1. 创建一个新的 Word 文档
    doc = Document()
    doc.add_heading('专利质检报告', level=1)

    # 2. 写入原始文档内容（可选）
    doc.add_heading('原始文档内容', level=2)
    doc.add_paragraph(parsed_content[:2000] + '...')  # 只展示前2000字

    # 3. 写入质检结果
    doc.add_heading('质检发现的问题', level=2)
    issues = result_json.get('issues', [])
    if issues:
        for issue in issues:
            p = doc.add_paragraph()
            p.add_run(f"【{issue.get('severity','提示')}】").bold = True
            p.add_run(f" {issue.get('description','')}\n")
            p.add_run(f"建议：{issue.get('suggestion','')}")
    else:
        doc.add_paragraph(result_json.get('raw_output', '未发现问题或问题无法解析。'))

    # 4. 保存文档
    doc.save(report_path)
    return report_path